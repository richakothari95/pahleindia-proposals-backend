"""
Word document generation service.
Populates Word_Template_new.docx with AI-generated proposal content.
"""

import os
import copy
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from models.generation import ProposalContent

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "Word_Template_new.docx")


def _replace_in_textboxes(doc: Document, replacements: dict[str, str]):
    """Replace placeholder text inside text boxes (txbxContent nodes) in the document body."""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    body = doc.element.body

    for txbx in body.iter(qn("w:txbxContent")):
        for t_elem in txbx.iter(qn("w:t")):
            for placeholder, replacement in replacements.items():
                if placeholder in (t_elem.text or ""):
                    t_elem.text = t_elem.text.replace(placeholder, replacement)


def _remove_table(doc: Document, table_index: int):
    """Remove a table from the document body by index."""
    if table_index < len(doc.tables):
        table = doc.tables[table_index]
        parent = table._element.getparent()
        if parent is not None:
            parent.remove(table._element)


def _add_heading(doc: Document, text: str, level: int):
    """Add a heading paragraph using PIF named styles if available, else default."""
    style_map = {
        1: "Heading1PahleIndia",
        2: "Heading2PahleIndia",
    }
    style_name = style_map.get(level, f"Heading {level}")
    # Fall back to built-in heading if custom style doesn't exist
    style_names = [s.name for s in doc.styles]
    if style_name not in style_names:
        style_name = f"Heading {level}"
    doc.add_heading(text, level=level)
    # Apply PIF style if available
    para = doc.paragraphs[-1]
    if style_map.get(level) in style_names:
        para.style = doc.styles[style_map[level]]


def _add_body(doc: Document, text: str):
    """Add a body text paragraph."""
    style_names = [s.name for s in doc.styles]
    style = "BodyTextPahleIndia" if "BodyTextPahleIndia" in style_names else "Normal"
    para = doc.add_paragraph(text)
    para.style = doc.styles[style]


def _add_bullet(doc: Document, text: str):
    """Add a bullet point paragraph."""
    style_names = [s.name for s in doc.styles]
    style = "BulletPOINT1StylePahleIndia" if "BulletPOINT1StylePahleIndia" in style_names else "List Bullet"
    para = doc.add_paragraph(text)
    try:
        para.style = doc.styles[style]
    except Exception:
        pass


def _add_numbered(doc: Document, text: str):
    """Add a numbered list paragraph."""
    style_names = [s.name for s in doc.styles]
    style = "NumberedListPahleIndia" if "NumberedListPahleIndia" in style_names else "List Number"
    para = doc.add_paragraph(text)
    try:
        para.style = doc.styles[style]
    except Exception:
        pass


class DocxService:
    def generate(self, content: ProposalContent) -> bytes:
        doc = Document(TEMPLATE_PATH)

        # Step 1: Replace cover page placeholders in text boxes
        _replace_in_textboxes(doc, {
            "<TITLE OF REPORT>": content.title,
            "<AUTHOR(S)/ CONTRIBUTORS>": content.authors,
            "<DESIGNATION>": "Research Team, Pahle India Foundation",
        })

        # Step 2: Remove the reference structure table (Table 0)
        _remove_table(doc, 0)

        # Step 3: Table of Contents
        _add_heading(doc, "Table of Contents", level=1)
        _add_body(doc, "[Right-click this line in Microsoft Word and select 'Update Field' to refresh the Table of Contents]")
        doc.add_page_break()

        # Step 4: Glossary
        if content.glossary:
            _add_heading(doc, "Glossary", level=1)
            for item in content.glossary:
                _add_body(doc, f"{item.term}: {item.definition}")
            doc.add_page_break()

        # Step 5: Executive Summary
        _add_heading(doc, "Executive Summary", level=1)
        _add_body(doc, content.executive_summary)
        doc.add_page_break()

        # Step 6: Problem Statement
        _add_heading(doc, "Problem Statement", level=1)
        _add_body(doc, content.problem_statement)

        # Step 7: Policy Context
        _add_heading(doc, "Policy Context", level=1)
        _add_body(doc, content.policy_context)

        # Step 8: Objectives
        _add_heading(doc, "Objectives", level=1)
        for obj in content.objectives:
            _add_numbered(doc, f"{obj.number}. {obj.text}")

        # Step 9: Methodology
        _add_heading(doc, "Methodology", level=1)
        _add_body(doc, content.methodology)

        # Step 10: Results
        _add_heading(doc, "Results and Findings", level=1)
        for result in content.results:
            _add_heading(doc, result.heading, level=2)
            _add_body(doc, result.content)
            if result.data_point:
                _add_body(doc, f"Key Data Point: {result.data_point}")

        # Step 11: Takeaways
        _add_heading(doc, "Takeaways and Recommendations", level=1)
        for t in content.takeaways:
            _add_bullet(doc, f"{t.number}. [{t.actor}] {t.recommendation}")

        # Step 12: Conclusion
        _add_heading(doc, "Conclusion", level=1)
        _add_body(doc, content.conclusion)

        # Step 13: Annexures
        if content.annexures:
            _add_heading(doc, "Annexures", level=1)
            _add_body(doc, content.annexures)

        # Step 14: References
        if content.sources:
            _add_heading(doc, "References", level=1)
            for i, src in enumerate(content.sources, 1):
                ref_text = f"{i}. {src.title}"
                if src.url:
                    ref_text += f". Available at: {src.url}"
                if src.date:
                    ref_text += f" (accessed {src.date})"
                _add_body(doc, ref_text)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
